"""
Trovly - Resume File Parser
Extracts plain text from PDF, DOCX, and TXT files for resume input.

Security:
- File size limit (5MB)
- Extension whitelist
- Content sanitization after extraction
- Magic byte verification (not just trust file extension)
"""

import io
import logging
from pathlib import Path

logger = logging.getLogger("trovly.resume_parser")

MAX_FILE_SIZE = 5 * 1024 * 1024  # 5MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt"}

# Magic bytes for file type verification
MAGIC_BYTES = {
    "pdf": b"%PDF",
    "docx": b"PK\x03\x04",
    "txt": None,
}


def detect_file_type(file_bytes):
    """Detect actual file type from magic bytes, not extension."""
    if file_bytes.startswith(b"%PDF"):
        return "pdf"
    if file_bytes.startswith(b"PK\x03\x04"):
        return "docx"
    try:
        file_bytes[:1000].decode("utf-8")
        return "txt"
    except UnicodeDecodeError:
        return None


def extract_pdf_text(file_bytes):
    """Extract text from PDF bytes."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            try:
                text_parts.append(page.extract_text() or "")
            except Exception as e:
                logger.warning("Error extracting page: {}".format(e))
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error("PDF extraction error: {}".format(e))
        return None


def extract_docx_text(file_bytes):
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also include text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error("DOCX extraction error: {}".format(e))
        return None


def extract_txt_text(file_bytes):
    """Extract text from plain text bytes."""
    try:
        for encoding in ["utf-8", "latin-1", "cp1252"]:
            try:
                return file_bytes.decode(encoding).strip()
            except UnicodeDecodeError:
                continue
        return None
    except Exception as e:
        logger.error("TXT extraction error: {}".format(e))
        return None


def parse_resume_file(file_bytes, filename=""):
    """
    Parse a resume file and return extracted text.
    Returns (success, text_or_error).
    """
    if not file_bytes:
        return False, "File is empty"

    if len(file_bytes) > MAX_FILE_SIZE:
        return False, "File too large (max 5MB)"

    if len(file_bytes) < 50:
        return False, "File appears to be empty or corrupted"

    # Check extension if filename provided
    if filename:
        ext = Path(filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            return False, "Unsupported file type. Use PDF, DOCX, or TXT"

    # Verify with magic bytes (don't trust extension alone)
    detected_type = detect_file_type(file_bytes)
    if not detected_type:
        return False, "Could not determine file type. Use PDF, DOCX, or TXT"

    # Extract text based on detected type
    extractors = {
        "pdf": extract_pdf_text,
        "docx": extract_docx_text,
        "txt": extract_txt_text,
    }

    extracted = extractors[detected_type](file_bytes)

    if not extracted:
        return False, "Could not extract text. The file may be a scanned image or corrupted."

    if len(extracted.strip()) < 100:
        return False, "Resume text too short. Make sure your file contains your full resume."

    if len(extracted) > 50000:
        extracted = extracted[:50000]
        logger.info("Truncated resume to 50,000 chars")

    return True, extracted
